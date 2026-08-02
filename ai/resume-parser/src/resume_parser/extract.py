"""Turning uploaded bytes into text (ADR-0016).

**This is the only module allowed to import a document library.** Everything else in the parser
works on `ExtractedText`, which is what keeps a swap from `pypdf` to `pdfplumber` a one-file change
rather than a rewrite. A bare ``import pypdf`` anywhere else means the port was bypassed and
ADR-0016's reversal argument is void.

Two rules shape this module more than extraction quality does:

**An uploaded document is hostile input** (``docs/features/resume-parsing.md``). A malformed file
must produce a refusal, never a traceback that escapes as a 500 — so every library call is wrapped
and every failure becomes an ``ExtractedText`` the pipeline can report honestly.

**Résumé text never appears in a log, an error, or a fixture** (``docs/architecture/privacy.md``).
No exception message here interpolates document content. That is easy to write correctly once and
easy to break during a debugging session, which is why it is stated at the top of the file.
"""

from __future__ import annotations

import io
import zipfile

import docx
import pypdf

from resume_parser.ports import ExtractedText

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PLAIN_TEXT_CONTENT_TYPE = "text/plain"

SUPPORTED_CONTENT_TYPES = frozenset({PDF_CONTENT_TYPE, DOCX_CONTENT_TYPE, PLAIN_TEXT_CONTENT_TYPE})


class UnsupportedDocumentError(Exception):
    """The content type is not one we read.

    Deliberately carries the content type and nothing else — never a filename, which is user data,
    and never a byte of the document.
    """

    def __init__(self, content_type: str) -> None:
        super().__init__(f"unsupported content type: {content_type!r}")
        self.content_type = content_type


def _page_has_image(page: pypdf.PageObject) -> bool:
    """Whether a page draws an image.

    This is what makes ``image_only`` reachable. A scanned résumé and a genuinely blank page both
    extract to an empty string, and telling a user "no text could be read" when the honest answer is
    "this is a scan, we cannot OCR it" wastes their time on the wrong fix.
    """
    try:
        resources = page.get("/Resources")
        if resources is None:
            return False
        xobjects = resources.get_object().get("/XObject")  # type: ignore[union-attr]
        if xobjects is None:
            return False
        for reference in xobjects.get_object().values():  # type: ignore[union-attr]
            if reference.get_object().get("/Subtype") == "/Image":
                return True
    except Exception:
        return False
    return False


def _extract_pdf(content: bytes) -> ExtractedText:
    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
    except Exception:
        return ExtractedText(text="", degraded_sections=("the whole document",))

    if reader.is_encrypted:
        # A password-protected résumé is a refusal, not a failure to try harder.
        return ExtractedText(
            text="", degraded_sections=("the whole document (password protected)",)
        )

    pieces: list[str] = []
    degraded: list[str] = []
    saw_image = False

    for number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if not page_text.strip():
            # A page that yielded nothing is named rather than dropped: this is what turns a
            # half-read document into `partial` instead of a confident half-profile.
            degraded.append(f"page {number}")
            if _page_has_image(page):
                saw_image = True
        pieces.append(page_text)

    text = "\n".join(pieces)
    # Only image-only if *nothing* was readable anywhere. A text résumé with one scanned page is a
    # partial read, not a scan.
    image_only = saw_image and not text.strip()
    return ExtractedText(
        text=text,
        degraded_sections=tuple(degraded) if text.strip() else (),
        image_only=image_only,
    )


def _extract_docx(content: bytes) -> ExtractedText:
    try:
        document = docx.Document(io.BytesIO(content))
    except (zipfile.BadZipFile, KeyError, ValueError, OSError):
        return ExtractedText(text="", degraded_sections=("the whole document",))

    lines = [paragraph.text for paragraph in document.paragraphs]

    # Tables are where two-column résumés hide their entire skills section. Read left-to-right,
    # row by row, and say so — a table flattened this way loses its column meaning, and a reader
    # who does not know that will trust the ordering.
    degraded: list[str] = []
    for table in document.tables:
        degraded.append("a table")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" ".join(cells))

    return ExtractedText(text="\n".join(lines), degraded_sections=tuple(dict.fromkeys(degraded)))


def extract(content: bytes, content_type: str) -> ExtractedText:
    """Read a document. Never raises on malformed input — only on a type we do not support.

    The distinction matters: an unsupported type is a caller error the gateway should have caught,
    while a corrupt PDF is an ordinary Tuesday that the user must be told about in words they can
    act on.
    """
    if content_type == PDF_CONTENT_TYPE:
        return _extract_pdf(content)
    if content_type == DOCX_CONTENT_TYPE:
        return _extract_docx(content)
    if content_type == PLAIN_TEXT_CONTENT_TYPE:
        return ExtractedText(text=content.decode("utf-8", errors="replace"))
    raise UnsupportedDocumentError(content_type)


class DocumentTextExtractor:
    """The `TextExtractor` port, implemented with `pypdf` and `python-docx` (ADR-0016)."""

    def extract(self, content: bytes, content_type: str) -> ExtractedText:
        return extract(content, content_type)
