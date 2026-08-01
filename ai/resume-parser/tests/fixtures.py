"""Synthetic documents for the extraction tests.

**Never a real person's résumé, and never real résumé text** (``docs/development/testing.md``,
``docs/architecture/privacy.md``). Everything here is built byte by byte in this file, so there is
no fixture file on disk that could quietly become someone's actual CV.

PDFs are hand-assembled rather than produced by a writer library. That keeps the test dependency
surface at zero and, more usefully, makes the *shape* of what is being tested explicit: a page with
a text operator, a page with only an image XObject, a file that is not a PDF at all.
"""

from __future__ import annotations

import io

import docx


def _pdf(objects: list[bytes]) -> bytes:
    """Assemble numbered objects into a valid PDF with a correct xref table."""
    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(out.tell())
        out.write(f"{number} 0 obj\n".encode())
        out.write(body)
        out.write(b"\nendobj\n")

    xref_at = out.tell()
    count = len(objects) + 1
    out.write(f"xref\n0 {count}\n".encode())
    out.write(b"0000000000 65535 f \n")
    for offset in offsets:
        out.write(f"{offset:010d} 00000 n \n".encode())
    out.write(f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF\n".encode())
    return out.getvalue()


def pdf_with_text(lines: list[str]) -> bytes:
    """A single-page PDF whose content stream draws the given lines."""
    drawn = "\n".join(
        f"BT /F1 12 Tf 72 {720 - index * 18} Td ({line}) Tj ET" for index, line in enumerate(lines)
    )
    stream = drawn.encode("latin-1", errors="replace")
    return _pdf(
        [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        ]
    )


def pdf_image_only() -> bytes:
    """A page that draws an image and no text — a scanned résumé, structurally."""
    image = bytes([0xFF, 0xD8, 0xFF, 0xD9])  # not a real JPEG; nothing decodes it here
    stream = b"q 200 0 0 200 100 500 cm /Im1 Do Q"
    return _pdf(
        [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /XObject << /Im1 5 0 R >> >> /Contents 4 0 R >>",
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
            b"<< /Type /XObject /Subtype /Image /Width 8 /Height 8 /ColorSpace /DeviceGray "
            b"/BitsPerComponent 8 /Filter /DCTDecode /Length "
            + str(len(image)).encode()
            + b" >>\nstream\n"
            + image
            + b"\nendstream",
        ]
    )


def pdf_blank() -> bytes:
    """A page with neither text nor image — genuinely empty, not a scan."""
    return _pdf(
        [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << >> >>",
        ]
    )


def not_a_document() -> bytes:
    """Bytes that are not any document format — the hostile-input case."""
    return b"\x00\x01\x02 this is not a PDF or a DOCX \xff\xfe"


def docx_with(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """A real .docx, built with python-docx so the file is genuinely well-formed."""
    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        for row in table_rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
